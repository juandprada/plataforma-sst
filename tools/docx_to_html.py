"""Convierte un .docx (o .docm) del SG-SST en el CUERPO HTML de una plantilla.

- Párrafos (con negrita e listas) -> <p> / <ul><li>.
- Tablas con celdas combinadas (gridSpan / vMerge) -> <table class="doc-tabla">.
- Tokeniza literales de la empresa-muestra -> {{EMPRESA}}, {{REPRESENTANTE_LEGAL}}, etc.
- NO incluye el encabezado (logo/código/versión); eso lo agrega la app.

Uso:
    python tools/docx_to_html.py --src "../1.PLANEAR/6.OBJETIVOS.docx" \
        --out plantillas/objetivos.html
"""
from __future__ import annotations

import argparse
import html
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

import docx
from docx.oxml.ns import qn

# Empresas-muestra usadas al crear las plantillas -> {{EMPRESA}}.
_EMPRESAS_MUESTRA = [
    "CK COMERCIALIZADORA UN MUNDO DE OPORTUNIDADES SAS",
    "DISTRIBUIDORA CONSUMAZ SAS",
    "CAUCHOS BARBERENA SAS",
    "MEDALLERIA DEPORTIVA DEL VALLE",
    "GAF TECHNOLOGY SOLUTIONS SAS",
]
# Representantes legales de muestra -> {{REPRESENTANTE_LEGAL}} (todos los del
# EMPRESAS.xlsx original, por si alguna plantilla usó otra empresa de base).
_REPS_MUESTRA = [
    "VICENTE CARLO CHAGUENDO REYES",
    "LIGNEY MOSQUERA CASTILLO",
    "LILIANA MARÍA OROZCO BERMÚDEZ",
    "NORBERTO SILVA VALENCIA",
    "JUAN CARLOS SOTO CORREA",
    "CARLOS VICENTE CHAGUENDO ECHEVERRY",
    "BRIAN ARMITAGE SALAZAR",
    "LEIDY VANESSA DUQUE CEBALLOS",
    "STIVEN FERNANDO MORA FRANCO",
    "MICHELLE ANDREA BENSOUR GONZÁLEZ",
    "FERNANDO ANDRES PARRA BERNAL",
    "JESUS ALBERTO ZULUAGA AGUIRRE",
]


def _variantes(nombre: str) -> list[str]:
    """Devuelve variantes de mayúsculas para cubrir el texto tal cual aparece."""
    return list(dict.fromkeys([nombre, nombre.title(), nombre.capitalize()]))


REEMPLAZOS: dict[str, str] = {}
for _e in _EMPRESAS_MUESTRA:
    for _v in _variantes(_e):
        REEMPLAZOS[_v] = "{{EMPRESA}}"
for _r in _REPS_MUESTRA:
    for _v in _variantes(_r):
        REEMPLAZOS[_v] = "{{REPRESENTANTE_LEGAL}}"
REEMPLAZOS["empresa GAF"] = "empresa {{EMPRESA}}"
REEMPLAZOS["ACTIVIDADES FINANCIERAS"] = "{{ACTIVIDAD}}"
# Ordena por longitud desc: reemplaza los literales largos antes que los cortos.
REEMPLAZOS = dict(sorted(REEMPLAZOS.items(), key=lambda kv: -len(kv[0])))


def tokenizar(texto: str) -> str:
    for lit, tok in REEMPLAZOS.items():
        texto = texto.replace(lit, tok)
    # Pares de años consecutivos (p.ej. "2024 y 2025") -> año actual y el siguiente.
    texto = re.sub(
        r"\b(\d{4}) y (\d{4})\b",
        lambda m: "{{ANIO}} y {{ANIO_SIGUIENTE}}"
        if int(m.group(2)) == int(m.group(1)) + 1
        else m.group(0),
        texto,
    )
    return texto


def esc(texto: str) -> str:
    # Escapa, conserva tokens {{...}} y pasa saltos de línea a <br>.
    escapado = html.escape(tokenizar(texto), quote=False)
    return escapado.replace("\n", "<br>")


def cargar_documento(path: Path) -> docx.document.Document:
    """Abre .docx o .docm (re-empaquetando el .docm como .docx)."""
    if path.suffix.lower() != ".docm":
        return docx.Document(str(path))
    tmp = Path(tempfile.mkdtemp()) / "conv.docx"
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.namelist():
            data = zin.read(item)
            if item == "[Content_Types].xml":
                data = data.replace(
                    b"application/vnd.ms-word.document.macroEnabled.main+xml",
                    b"application/vnd.openxmlformats-officedocument."
                    b"wordprocessingml.document.main+xml",
                )
            zout.writestr(item, data)
    return docx.Document(str(tmp))


def _tiene_salto_pagina(p_el) -> bool:
    """True si el párrafo fuerza un salto de página (w:br type=page o sectPr)."""
    for br in p_el.iter(qn("w:br")):
        if br.get(qn("w:type")) == "page":
            return True
    pPr = p_el.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
        return True  # fin de sección => nueva página (salvo la sección final del body)
    return False


def parrafo_html(p) -> str | None:
    txt = p.text.strip()
    if not txt:
        return None
    runs = [r for r in p.runs if r.text.strip()]
    todo_negrita = bool(runs) and all(r.bold for r in runs)
    es_lista = p._p.find(qn("w:pPr")) is not None and (
        p._p.find(qn("w:pPr")).find(qn("w:numPr")) is not None
    )
    cont = esc(txt)
    if es_lista:
        return f"<li>{cont}</li>"
    if todo_negrita:
        return f'<p class="doc-h">{cont}</p>'
    return f"<p>{cont}</p>"


def _grid_span(tc) -> int:
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        gs = tcPr.find(qn("w:gridSpan"))
        if gs is not None:
            return int(gs.get(qn("w:val")))
    return 1


def _vmerge(tc):
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        return None
    val = vm.get(qn("w:val"))
    return "restart" if val == "restart" else "continue"


def _hex(val: str | None) -> str | None:
    """Normaliza un color OOXML a #RRGGBB; ignora 'auto'/vacío."""
    if not val or val.lower() == "auto":
        return None
    val = val.strip().lstrip("#")
    return "#" + val.upper() if len(val) == 6 else None


def _fill(tc) -> str | None:
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    s = tcPr.find(qn("w:shd"))
    return _hex(s.get(qn("w:fill"))) if s is not None else None


def _text_color(tc) -> str | None:
    for r in tc.iter(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            c = rPr.find(qn("w:color"))
            if c is not None:
                h = _hex(c.get(qn("w:val")))
                if h:
                    return h
    return None


def _para_text(p_el) -> str:
    """Texto de un <w:p>: une runs sin salto; w:br/w:cr -> salto de línea."""
    partes = []
    for node in p_el.iter():
        if node.tag == qn("w:t"):
            partes.append(node.text or "")
        elif node.tag in (qn("w:br"), qn("w:cr")):
            partes.append("\n")
    return "".join(partes)


def _cell_text(tc) -> str:
    """Texto de una celda: párrafos separados por salto (no cada run)."""
    return "\n".join(_para_text(p) for p in tc.findall(qn("w:p"))).strip()


def _row_height_px(tr) -> int | None:
    """Altura de fila (w:trHeight, en twips) convertida a px, o None."""
    trPr = tr.find(qn("w:trPr"))
    if trPr is None:
        return None
    th = trPr.find(qn("w:trHeight"))
    if th is None:
        return None
    val = th.get(qn("w:val"))
    if not val or not val.isdigit():
        return None
    return round(int(val) / 15)  # 1440 twips/pulgada, 96 px/pulgada


_RE_FIRMA = re.compile(r"^_{5,}")


def tabla_html(tab) -> str:
    tbl = tab._tbl
    filas = tbl.findall(qn("w:tr"))
    abiertos = {}  # col -> celda origen (para rowspan de vMerge)
    grid = []
    alturas = []
    for tr in filas:
        alturas.append(_row_height_px(tr))
        col = 0
        celdas = []
        for tc in tr.findall(qn("w:tc")):
            span = _grid_span(tc)
            vm = _vmerge(tc)
            texto = _cell_text(tc)
            if vm == "continue" and col in abiertos:
                abiertos[col]["rowspan"] += 1
                col += span
                continue
            # ¿negrita? true si hay algún run en negrita con texto.
            bold = any(
                (b := r.find(qn("w:rPr"))) is not None
                and b.find(qn("w:b")) is not None
                for r in tc.iter(qn("w:r"))
                if (r.find(qn("w:t")) is not None)
            )
            celda = {
                "texto": texto,
                "colspan": span,
                "rowspan": 1,
                "col": col,
                "bold": bold,
                "fill": _fill(tc),
                "color": _text_color(tc),
                "firma": bool(_RE_FIRMA.match(texto)),
            }
            celdas.append(celda)
            if vm == "restart":
                abiertos[col] = celda
            else:
                abiertos.pop(col, None)
            col += span
        grid.append(celdas)

    # ¿Primera fila es encabezado? (tabla con varias filas y fila0 en negrita)
    header = (
        len(grid) > 2
        and grid[0]
        and all(c["bold"] for c in grid[0] if c["texto"])
        and any(c["texto"] for c in grid[0])
    )

    out = ['<table class="doc-tabla">']
    for i, fila in enumerate(grid):
        out.append(f'<tr style="height:{alturas[i]}px">' if alturas[i] else "<tr>")
        tag = "th" if (header and i == 0) else "td"
        for c in fila:
            attrs = ""
            if c["colspan"] > 1:
                attrs += f' colspan="{c["colspan"]}"'
            if c["rowspan"] > 1:
                attrs += f' rowspan="{c["rowspan"]}"'
            if c["firma"]:
                attrs += ' class="firma-cel"'
            # Respeta el sombreado y color de texto originales de la celda.
            estilos = []
            if c["fill"]:
                estilos.append(f"background-color:{c['fill']}")
            if c["color"]:
                estilos.append(f"color:{c['color']}")
            if estilos:
                attrs += f' style="{";".join(estilos)}"'
            texto = c["texto"]
            # Firma automática de la consultora: en celdas con "Karen" junto a una
            # línea de firma, inserta {{FIRMA_CONSULTORA}} sobre esa línea.
            if texto and "karen" in texto.lower() and re.search(
                r"(?mi)^[ \t]*(?:firma[ \t]*)?_{5,}", texto
            ):
                texto = re.sub(
                    r"(?mi)^([ \t]*(?:firma[ \t]*)?_{5,})",
                    "@@FIRMACONS@@\\1",
                    texto,
                    count=1,
                )
            contenido = esc(texto) if texto else "&nbsp;"
            contenido = contenido.replace("@@FIRMACONS@@", "{{FIRMA_CONSULTORA}}<br>")
            if c["bold"] and tag == "td" and texto:
                contenido = f"<strong>{contenido}</strong>"
            out.append(f"<{tag}{attrs}>{contenido}</{tag}>")
        out.append("</tr>")
    out.append("</table>")
    return "\n".join(out)


def convertir(path: Path) -> str:
    d = cargar_documento(path)
    partes: list[str] = []
    lista_abierta = False
    para_i = tab_i = 0
    for child in d.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = d.paragraphs[para_i]
            para_i += 1
            frag = parrafo_html(p)
            salto = _tiene_salto_pagina(p._p)
            if frag is not None:
                if frag.startswith("<li>"):
                    if not lista_abierta:
                        partes.append("<ul>")
                        lista_abierta = True
                    partes.append(frag)
                else:
                    if lista_abierta:
                        partes.append("</ul>")
                        lista_abierta = False
                    partes.append(frag)
            if salto:
                if lista_abierta:
                    partes.append("</ul>")
                    lista_abierta = False
                partes.append('<div class="salto-pagina"></div>')
        elif child.tag == qn("w:tbl"):
            if lista_abierta:
                partes.append("</ul>")
                lista_abierta = False
            partes.append(tabla_html(d.tables[tab_i]))
            tab_i += 1
    if lista_abierta:
        partes.append("</ul>")
    return "\n".join(partes) + "\n"


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--src", required=True)
    ap.add_argument("--out", required=True)
    args = ap.parse_args()
    html_body = convertir(Path(args.src))
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        f"<!-- Generado desde {Path(args.src).name} por docx_to_html.py; "
        f"revisar tokens y ajustes. -->\n{html_body}",
        encoding="utf-8",
    )
    tokens = sorted(set(re.findall(r"\{\{([A-Z0-9_]+)\}\}", html_body)))
    print(f"OK -> {out}  tokens: {tokens}")


if __name__ == "__main__":
    main()
